# mdr_7stage_app.py
"""
MDR 7-Stage Cleaner (Local PDFs/DOCX/TXT)

What's new (per user requests):
- Automatic per-page OCR (no toggle). If a page looks image-only/sparse, we OCR it.
- Full scanned-PDF detection: ask confirmation to run full-document OCR (with progress).
- No silent drops: show an explicit warning when text is removed at any stage.
- JSONL always includes tables and figure OCR (buttons never disabled).

Also includes:
- Stage checkboxes, progress bars, skip-on-fail prompts
- Language gate modes (Lenient/Strict/Off, math/code-aware)
- Book boilerplate cleanup, repetition removal, fuzzy/exact dedup
- Math & code -> LaTeX transform
- Per-file & combined downloads: TXT, JSONL, CSV, DOCX

Install (minimum):
  pip install streamlit pymupdf docx2txt pytesseract pillow datasketch fasttext pandas
Optional:
  pip install python-docx camelot-py[cv]
"""

import io, os, re, json, uuid, tempfile, hashlib
from hashlib import sha256
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

import streamlit as st
from PIL import Image
import fitz  # PyMuPDF
import docx2txt
import pytesseract
from datasketch import MinHash, MinHashLSH
import pandas as pd

# Optional deps
try:
    import camelot
    HAS_CAMELOT = True
except Exception:
    HAS_CAMELOT = False

try:
    from docx import Document
    HAS_PYDOCX = True
except Exception:
    HAS_PYDOCX = False

# ---------------- App Config ----------------
st.set_page_config(page_title="MDR 7-Stage Cleaner", layout="wide")
st.title("MDR 7-Stage Cleaner (Local Files)")
st.caption("PDF, DOCX, and TXT → cleaned text (math/code → LaTeX) → TXT/JSONL/CSV/DOCX (per-file & combined)")

# ---------------- Pipeline Config ----------------
TESSERACT_CMD = os.environ.get("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
if os.name == "nt":
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

LANG_MODEL_PATH = os.environ.get("FASTTEXT_LID_PATH", "lid.176.bin")
LANG_THRESHOLD = 0.65
MIN_DOC_WORDS = 50
MIN_SENT_LEN = 20

STAGE_LABELS = {
    1: "1) Text Extraction",
    2: "2) Language Identification",
    3: "3) Document-wise Filtering",
    4: "4) Line-wise Filtering (incl. book boilerplate)",
    5: "5) In-document Repetition Removal",
    6: "6) Fuzzy Deduplication (MinHash)",
    7: "7) Exact Deduplication (span hashes)",
    8: "8) Math & Code → LaTeX (transform)",
}
STAGE_HELP = {
    1: "Parse text from local PDFs/DOCX/TXT. OCR is automatic for image-only pages; full scanned PDFs prompt for full OCR.",
    2: "Keep English only (Lenient/Strict/Off). Lenient is math/code-aware.",
    3: "Heuristics: min length, symbol ratio, whitespace cleanup.",
    4: "Remove headers/footers, page numbers, start/end markers.",
    5: "Remove repeated sentences within the document.",
    6: "Block near-dupes vs. prior docs/chunks using LSH.",
    7: "Block exact repeated spans across docs/chunks.",
    8: "Convert math symbols & code blocks to LaTeX (inline & block).",
}

UNWANTED_PATTERNS = [
    "subscribe","follow us","click here","share on","advertisement",
    "back to top","login","sign up","terms of service","copyright",
    "all rights reserved"
]
BOOK_EXTRA_PATTERNS = [
    r'^\s*\d+\s*$', r'^\s*page\s*\d+\s*$', r'^\s*\d+\s*/\s*\d+\s*$',
    "chapter","contents","table of contents","bibliography","references",
    "index","start of the book","end of the book","project gutenberg",
    "all rights reserved",
]

# Global dedup
LSH = MinHashLSH(threshold=0.8, num_perm=128)
SEEN_SPAN_HASHES: set[str] = set()

# ---------------- Utilities ----------------
def sha(text: str) -> str:
    return sha256(text.encode("utf-8", errors="ignore")).hexdigest()

def norm_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()

def simple_sentences(text: str) -> List[str]:
    return [p.strip() for p in re.split(r"(?<=[.!?])\s+", text) if p.strip()]

@st.cache_resource(show_spinner=False)
def load_lang_model():
    import fasttext
    return fasttext.load_model(LANG_MODEL_PATH)

# NEW: fingerprint helper so outputs aren't reset on every rerun
def _fingerprint_uploads(files) -> list[tuple[str,int,str]]:
    fps = []
    for f in files:
        b = f.getvalue()  # bytes; doesn't consume the file
        h = sha256(b).hexdigest()[:16]
        fps.append((f.name, len(b), h))
    fps.sort()
    return fps

# ---- Language gate (math/code-aware) ----
MATH_OR_CODE_TOKENS = set([
    "int", "sum", "sqrt", "lim", "sin", "cos", "tan", "log",
    "def", "class", "import", "return", "for", "while", "if", "else"
])
MATH_OR_CODE_CHARS = set("∫∑√≤≥≠≈±×÷→←∞πΔαβγθλμσφωΩ{}[]()=+−-*/^_|<>`#;$\\:")

def looks_like_math_or_code(text: str) -> bool:
    if any(ch in text for ch in MATH_OR_CODE_CHARS):
        return True
    ascii_ratio = sum(1 for c in text if ord(c) < 128) / max(1, len(text))
    if ascii_ratio > 0.9 and any(tok in text for tok in MATH_OR_CODE_TOKENS):
        return True
    if "```" in text or re.search(r"`[^`]+`", text):
        return True
    return False

def en_gate(text: str, mode: str = "Lenient") -> bool:
    """Return True if should be treated as English."""
    t = norm_ws(text)
    if not t: return False
    if mode == "Off": return True

    idxs = [0, int(len(t)*0.25), int(len(t)*0.5), int(len(t)*0.75), max(0, len(t)-5000)]
    samples = [t[i:i+5000] for i in idxs]
    model = load_lang_model()

    votes_en = 0
    for s in samples:
        if not s: continue
        lab, prob = model.predict(s, k=1)
        if lab[0] == "__label__en" and float(prob[0]) >= (LANG_THRESHOLD if mode=="Strict" else 0.55):
            votes_en += 1

    if votes_en == 0 and looks_like_math_or_code(t):
        return True
    return votes_en >= (3 if mode == "Strict" else 1)

# ---------------- Extraction helpers ----------------
def read_txt(file) -> str:
    data = file.read()
    for enc in ("utf-8","latin-1"):
        try: return data.decode(enc)
        except Exception: pass
    return data.decode(errors="ignore")

def read_docx(file) -> Tuple[str, List[List[List[str]]]]:
    tmp = io.BytesIO(file.read())
    text = docx2txt.process(tmp) or ""
    tables: List[List[List[str]]] = []
    try:
        if HAS_PYDOCX:
            tmp.seek(0)
            doc = Document(tmp)
            for t in doc.tables:
                cells = [[norm_ws(c.text) for c in row.cells] for row in t.rows]
                tables.append(cells)
    except Exception:
        pass
    return text, tables

def table_to_markdown(cells: List[List[str]]) -> str:
    if not cells: return ""
    header = cells[0]
    body = cells[1:] if len(cells) > 1 else []
    def row(r): return "| " + " | ".join((c or "").replace("|","\\|") for c in r) + " |"
    md = [row(header), "| " + " | ".join(["---"]*len(header)) + " |"]
    for r in body: md.append(row(r))
    return "\n".join(md)

def extract_pdf_tables_via_camelot(path_like: str) -> List[str]:
    mds: List[str] = []
    if not HAS_CAMELOT: return mds
    try:
        tables = camelot.read_pdf(path_like, pages="all", flavor="lattice")
        for t in tables: mds.append(t.df.to_markdown(index=False))
    except Exception:
        try:
            tables = camelot.read_pdf(path_like, pages="all", flavor="stream")
            for t in tables: mds.append(t.df.to_markdown(index=False))
        except Exception:
            pass
    return mds

def selective_page_ocr(page) -> Optional[str]:
    """OCR a single page to text using a 2x raster."""
    try:
        pix = page.get_pixmap(matrix=fitz.Matrix(2,2))
        img = Image.frombytes("RGB",[pix.width,pix.height],pix.samples)
        ocr = norm_ws(pytesseract.image_to_string(img))
        return ocr if ocr else None
    except Exception:
        return None

def extract_pdf_auto_ocr(up_file) -> Tuple[str, List[str], List[Dict[str, Any]], bool, List[int]]:
    b = up_file.read()
    text_pages: List[str] = []
    tables_md: List[str] = []
    figures: List[Dict[str, Any]] = []
    image_only_pages: List[int] = []

    with fitz.open(stream=b, filetype="pdf") as doc:
        empty_or_sparse = 0
        for pno, page in enumerate(doc, start=1):
            txt = page.get_text("text") or ""
            if len(txt.strip()) < 80:
                empty_or_sparse += 1
                ocr = selective_page_ocr(page)
                if ocr:
                    figures.append({"page": pno, "figure_text": ocr})
                    image_only_pages.append(pno)
                    txt = txt + "\n" + ocr
            text_pages.append(txt)

        # Table extraction via camelot (needs temp path)
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(b); tmp.flush()
                tables_md = extract_pdf_tables_via_camelot(tmp.name)
            os.unlink(tmp.name)
        except Exception:
            pass

    text = "\n".join(text_pages)
    is_scanned_candidate = (empty_or_sparse >= max(2, int(0.6 * len(text_pages))))
    return text, tables_md, figures, is_scanned_candidate, image_only_pages

def full_document_ocr(up_file, total_pages_hint: Optional[int] = None) -> str:
    b = up_file.read()
    pages_text: List[str] = []
    with fitz.open(stream=b, filetype="pdf") as doc:
        total = total_pages_hint or len(doc)
        pbar = st.progress(0.0, text="Running full-document OCR…")
        for idx, page in enumerate(doc, start=1):
            ocr = selective_page_ocr(page) or ""
            pages_text.append(ocr)
            pbar.progress(idx / total, text=f"OCR {idx}/{total} pages")
    return "\n".join(pages_text)

# ---------------- Cleaning stages ----------------
def document_wise_filters(text: str) -> Optional[str]:
    if not text.strip():
        return None
    words = re.findall(r"\b\w+\b", text)
    if len(words) < MIN_DOC_WORDS:
        return None
    if len(re.findall(r"[^\w\s]", text)) / max(1,len(text)) > 0.25:
        return None
    lines = [ln.strip() for ln in text.splitlines() if len(ln.strip()) >= MIN_SENT_LEN]
    return norm_ws(" ".join(lines))

def line_wise_filter(text: str) -> str:
    kept = []
    for ln in text.splitlines():
        lnl = ln.lower().strip()
        if any(p in lnl for p in UNWANTED_PATTERNS): continue
        if re.fullmatch(r"\s*\d+\s*", lnl): continue
        if re.fullmatch(r"\s*page\s*\d+\s*", lnl): continue
        if re.fullmatch(r"\s*\d+\s*/\s*\d+\s*", lnl): continue
        if any(k in lnl for k in ["project gutenberg","start of the book","end of the book"]): continue
        if any(k in lnl for k in ["chapter","table of contents","contents","bibliography","references","index"]):
            if len(lnl) <= 40: continue
        kept.append(ln)
    return norm_ws(" ".join(kept))

def remove_in_document_repetition(text: str) -> str:
    sents, out, seen = simple_sentences(text), [], set()
    for s in sents:
        key = sha(s.lower())
        if key in seen: continue
        seen.add(key); out.append(s)
    return norm_ws(" ".join(out))

def is_fuzzy_duplicate(text: str, key: str) -> bool:
    m = MinHash(num_perm=128)
    for w in set(re.findall(r"\w+", text.lower())):
        m.update(w.encode("utf-8"))
    res = LSH.query(m)
    if res: return True
    LSH.insert(key, m); return False

def has_exact_dup_spans(text: str, span_words: int = 30) -> bool:
    words = re.findall(r"\w+", text)
    if len(words) < span_words: return False
    step = max(1, span_words // 2)
    for i in range(0, len(words) - span_words + 1, step):
        span = " ".join(words[i:i+span_words])
        if sha(span) in SEEN_SPAN_HASHES: return True
    for i in range(0, len(words) - span_words + 1, step):
        SEEN_SPAN_HASHES.add(sha(" ".join(words[i:i+span_words])))
    return False

# ---- Code & Math → LaTeX ----
CODE_FENCES = re.compile(r"```([a-zA-Z0-9_+-]*)\n(.*?)```", re.DOTALL)
INDENT_CODE = re.compile(r"(^\s{4,}.*(?:\n\s{4,}.*)*)", re.MULTILINE)
INLINE_CODE = re.compile(r"`([^`]+)`")

GREEK_MAP = {
    "α":"\\alpha","β":"\\beta","γ":"\\gamma","δ":"\\delta","ε":"\\epsilon","θ":"\\theta",
    "λ":"\\lambda","μ":"\\mu","π":"\\pi","σ":"\\sigma","φ":"\\phi","ω":"\\omega",
    "Ω":"\\Omega","Δ":"\\Delta"
}
SYMBOL_MAP = {
    "∫":"\\int","∑":"\\sum","√":"\\sqrt{}","≤":"\\leq","≥":"\\geq","≠":"\\neq","≈":"\\approx",
    "±":"\\pm","×":"\\times","÷":"\\div","→":"\\to","←":"\\leftarrow","∞":"\\infty"
}
SUPERSCRIPTS = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁽⁾","0123456789+-()")
SUBSCRIPTS   = str.maketrans("₀₁₂₃₄₅₆₇₈₉₊₋₍₎","0123456789+-()")

def to_latex_listings(code: str, language: str = "text") -> str:
    code = code.replace("\\","\\\\")
    return f"\\begin{{lstlisting}}[language={language}]\n{code}\n\\end{{lstlisting}}"

def convert_code_blocks(text: str) -> str:
    # fenced
    def repl_f(m):
        lang = m.group(1) or "text"
        body = m.group(2)
        return to_latex_listings(body, lang)
    text = CODE_FENCES.sub(repl_f, text)
    # indented
    def repl_i(m):
        body = re.sub(r"^\s{4}", "", m.group(1), flags=re.MULTILINE)
        return to_latex_listings(body, "text")
    text = INDENT_CODE.sub(repl_i, text)
    # inline backticks
    text = INLINE_CODE.sub(lambda m: f"\\texttt{{{m.group(1)}}}", text)
    return text

def replace_unicode_math(s: str) -> str:
    s = re.sub(r"([A-Za-z0-9])([⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻]+)", lambda m: f"{m.group(1)}^{{{m.group(2).translate(SUPERSCRIPTS)}}}", s)
    s = re.sub(r"([A-Za-z0-9])([₀₁₂₃₄₅₆₇₈₉₊₋]+)", lambda m: f"{m.group(1)}_{{{m.group(2).translate(SUBSCRIPTS)}}}", s)
    for u, lx in {**GREEK_MAP, **SYMBOL_MAP}.items():
        s = s.replace(u, lx)
    return s

MATH_WRAP_PATTERN = re.compile(r"(\\(int|sum|sqrt|leq|geq|neq|approx|pm|times|div|to|leftarrow|infty|alpha|beta|gamma|delta|epsilon|theta|lambda|mu|pi|sigma|phi|omega)\b[^$]*)")

def wrap_inline_math(s: str) -> str:
    def add_dollars(seg: str) -> str:
        seg = seg.strip()
        if seg.startswith("$") and seg.endswith("$"): return seg
        return f"${seg}$"
    return MATH_WRAP_PATTERN.sub(lambda m: add_dollars(m.group(1)), s)

def transform_code_and_math_to_latex(text: str) -> str:
    if not text: return text
    text = convert_code_blocks(text)
    text = replace_unicode_math(text)
    sents = simple_sentences(text)
    sents = [wrap_inline_math(s) for s in sents]
    return " ".join(sents)

def rag_chunks(text: str, max_tokens: int = 512) -> List[str]:
    max_chars = max_tokens * 4
    text = norm_ws(text)
    chunks = []
    i = 0
    while i < len(text):
        j = min(len(text), i + max_chars)
        cut = text.rfind(". ", i, j)
        j = cut + 1 if cut > i else j
        chunks.append(text[i:j].strip())
        i = j
    return [c for c in chunks if c]

# ---------------- Sidebar ----------------
st.sidebar.header("Options")
auto_skip_fail = st.sidebar.checkbox("Auto-skip failed stages", False)
reset_outputs = st.sidebar.checkbox("Reset previous results when new files are selected", True)
lang_mode = st.sidebar.selectbox("Language gate mode", ["Lenient","Strict","Off"], index=0,
                                 help="Lenient = sample-based + math/code-aware. Off = skip Stage 2.")

st.sidebar.header("Stage Selection")
selected_stages = []
for num in range(1, 9):
    label = STAGE_LABELS[num]
    help_text = STAGE_HELP[num]
    if st.sidebar.checkbox(label, value=True, help=help_text, key=f"stage_{num}"):
        selected_stages.append(num)

with st.sidebar.expander("What each stage does", expanded=False):
    for n in range(1, 9):
        st.markdown(f"**{STAGE_LABELS[n]}** — {STAGE_HELP[n]}")

uploaded = st.file_uploader("Upload PDF / DOCX / TXT (multiple)",
                            accept_multiple_files=True,
                            type=["pdf","docx","txt"])

# ---------------- NEW: Big processing trigger button ----------------
if "do_process" not in st.session_state:
    st.session_state.do_process = False

process_clicked = st.button(
    "▶️ Process files now",
    type="primary",
    use_container_width=True,
    disabled=not uploaded,
    key="process_btn",
)
if process_clicked:
    st.session_state.do_process = True

# ---------------- NEW: Persist results across reruns ----------------
if "last_upload_fps" not in st.session_state:
    st.session_state.last_upload_fps = []

uploads_changed = False
if uploaded:
    cur_fps = _fingerprint_uploads(uploaded)
    uploads_changed = (cur_fps != st.session_state.last_upload_fps)
    if uploads_changed and reset_outputs:
        # Only clear when the actual selection of files changed
        st.session_state.files = {}
        st.session_state.do_process = False
    st.session_state.last_upload_fps = cur_fps

# ---------------- Session stores ----------------
if "files" not in st.session_state:
    st.session_state.files: Dict[str, Dict[str, Any]] = {}

def ask_skip(stage_num: int, file_key: str) -> bool:
    c1, c2 = st.columns(2)
    skip = c1.button(f"Skip {STAGE_LABELS[stage_num]} and continue", key=f"skip_{stage_num}_{file_key}")
    stop = c2.button("Stop processing", key=f"stop_{stage_num}_{file_key}")
    if stop: st.stop()
    return skip

# ---------------- Main processing ----------------
# Only process when the big button was clicked
if uploaded and st.session_state.do_process:
    progress = st.progress(0.0)
    total = len(uploaded)

    for f_idx, up in enumerate(uploaded, 1):
        file_key = f"{f_idx}_{uuid.uuid4().hex[:6]}"
        st.subheader(f"Processing: {up.name}")

        ext = up.name.lower().rsplit(".", 1)[-1]
        source_id = f"{up.name}::{uuid.uuid4().hex[:8]}"

        text = ""
        tables_md: List[str] = []
        figures: List[Dict[str, Any]] = []

        # 1) Extraction (auto OCR where needed)
        if 1 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[1]}"):
                    if ext == "pdf":
                        text, tables_md, figures, is_scanned, image_ocr_pages = extract_pdf_auto_ocr(up)
                        if is_scanned:
                            st.warning("This PDF looks scanned (most pages image-only). "
                                       "You can run full-document OCR to extract more text (slower).")
                            if st.button(f"Run full-document OCR for {up.name}", key=f"full_ocr_btn_{file_key}"):
                                up.seek(0)
                                text = full_document_ocr(up)
                                st.success("Full-document OCR complete.")
                        else:
                            if image_ocr_pages:
                                st.info(f"Light OCR applied on pages: {', '.join(map(str, image_ocr_pages[:10]))}"
                                        + ("…" if len(image_ocr_pages) > 10 else ""))
                    elif ext == "docx":
                        text, docx_tables = read_docx(up)
                        tables_md = [table_to_markdown(t) for t in docx_tables if t]
                    elif ext == "txt":
                        text = read_txt(up)
                    else:
                        st.warning("Unsupported file type; skipping"); continue
                    text = norm_ws(text)
            except Exception as e:
                st.error(f"{STAGE_LABELS[1]} failed: {e}")
                if not (auto_skip_fail or ask_skip(1, file_key)): st.stop()

        st.text_area("Raw extracted text (first 2k)", text[:2000], height=180)

        dropped = False  # track if/where text is cleared

        # 2) Language ID (math/code aware)
        if text and 2 in selected_stages and lang_mode != "Off":
            try:
                with st.spinner(f"Running {STAGE_LABELS[2]} [{lang_mode}]"):
                    if not en_gate(text[:50_000], mode=lang_mode):
                        st.warning("Stage 2 removed text: language gate deemed non-English.")
                        text = ""; dropped = True
            except Exception as e:
                st.error(f"{STAGE_LABELS[2]} failed: {e}")
                if not (auto_skip_fail or ask_skip(2, file_key)): st.stop()

        # 3) Document-wise filters
        if text and 3 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[3]}"):
                    t2 = document_wise_filters(text)
                    if not t2:
                        st.warning("Stage 3 removed text: too short or too many symbols after cleanup.")
                        text = ""; dropped = True
                    else:
                        text = t2
            except Exception as e:
                st.error(f"{STAGE_LABELS[3]} failed: {e}")
                if not (auto_skip_fail or ask_skip(3, file_key)): st.stop()

        # 4) Line-wise filters
        if text and 4 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[4]}"):
                    t2 = line_wise_filter(text)
                    if not t2.strip():
                        st.warning("Stage 4 removed text: considered book boilerplate/navigation.")
                        text = ""; dropped = True
                    else:
                        text = t2
            except Exception as e:
                st.error(f"{STAGE_LABELS[4]} failed: {e}")
                if not (auto_skip_fail or ask_skip(4, file_key)): st.stop()

        # 5) In-document repetition
        if text and 5 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[5]}"):
                    t2 = remove_in_document_repetition(text)
                    if not t2.strip():
                        st.warning("Stage 5 removed text: all sentences considered duplicates.")
                        text = ""; dropped = True
                    else:
                        text = t2
            except Exception as e:
                st.error(f"{STAGE_LABELS[5]} failed: {e}")
                if not (auto_skip_fail or ask_skip(5, file_key)): st.stop()

        # 8) Math & Code → LaTeX
        if text and 8 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[8]}"):
                    text = transform_code_and_math_to_latex(text)
            except Exception as e:
                st.error(f"{STAGE_LABELS[8]} failed: {e}")
                if not (auto_skip_fail or ask_skip(8, file_key)): st.stop()

        # 6) Fuzzy dedup
        if text and 6 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[6]}"):
                    if is_fuzzy_duplicate(text, f"doc::{source_id}"):
                        st.warning("Stage 6 removed text: fuzzy duplicate of a previous document.")
                        text = ""; dropped = True
            except Exception as e:
                st.error(f"{STAGE_LABELS[6]} failed: {e}")
                if not (auto_skip_fail or ask_skip(6, file_key)): st.stop()

        # 7) Exact dedup
        if text and 7 in selected_stages:
            try:
                with st.spinner(f"Running {STAGE_LABELS[7]}"):
                    if has_exact_dup_spans(text):
                        st.warning("Stage 7 removed text: shares exact spans with prior documents.")
                        text = ""; dropped = True
            except Exception as e:
                st.error(f"{STAGE_LABELS[7]} failed: {e}")
                if not (auto_skip_fail or ask_skip(7, file_key)): st.stop()

        # Preview of cleaned text (warn explicitly if empty)
        st.text_area("Cleaned (preview, first 2k)", (text[:2000] if text else ""), height=220)
        if not text and not dropped:
            st.info("No cleaned text available (document may be image-only or was skipped). "
                    "Tables/OCR snippets are still exported.")

        # -------- Build per-file artifacts (JSONL always has tables/figures) --------
        records: List[Dict[str, Any]] = []

        # text chunks
        if text:
            for i, ch in enumerate(rag_chunks(text)):
                rid = f"{source_id}#text#{i}"
                if 6 in selected_stages and is_fuzzy_duplicate(ch, rid): continue
                if 7 in selected_stages and has_exact_dup_spans(ch): continue
                records.append({
                    "id": rid,
                    "source": up.name,
                    "chunk_type": "text",
                    "text": ch,
                    "meta": {"stage": "cleaned+latex" if 8 in selected_stages else "cleaned",
                             "filetype": ext, "index": i}
                })

        # tables (always include)
        for i, md in enumerate(tables_md):
            records.append({
                "id": f"{source_id}#table#{i}",
                "source": up.name,
                "chunk_type": "table",
                "table_markdown": md,
                "meta": {"stage":"extracted","filetype": ext, "index": i}
            })

        # figures OCR (always include)
        for i, fg in enumerate(figures):
            records.append({
                "id": f"{source_id}#figure#{i}",
                "source": up.name,
                "chunk_type": "figure",
                "figure_text": fg.get("figure_text",""),
                "meta": {"page": fg.get("page"), "stage": "ocr"}
            })

        st.session_state.files[up.name] = {
            "text": text,
            "tables": tables_md,
            "figures": figures,
            "records": records,
            "meta": {
                "filename": up.name,
                "language": "en" if text else "n/a",
                "word_count": len(re.findall(r"\b\w+\b", text)) if text else 0,
                "num_text_chunks": sum(1 for r in records if r["chunk_type"] == "text"),
                "num_tables": len(tables_md),
                "num_figures": len(figures),
                "processed_at": datetime.utcnow().isoformat()
            }
        }

        st.write("Added {} records: {}".format(
            len(records),
            ", ".join(sorted(set(r["chunk_type"] for r in records))) or "none"
        ))

        progress.progress(f_idx / total)

    # prevent re-processing on next rerun (downloads won’t retrigger cleaning)
    st.session_state.do_process = False

# ---------------- Downloads ----------------
if st.session_state.files:
    st.markdown("---")
    st.subheader("Downloads")

    # Combined JSONL
    all_records = []
    for v in st.session_state.files.values():
        all_records.extend(v["records"])
    jsonl = "\n".join(json.dumps(r, ensure_ascii=False) for r in all_records)
    st.download_button("📦 Download ALL (JSONL)", jsonl, file_name="all_records.jsonl")

    # Combined CSV (metadata)
    meta_df = pd.DataFrame([v["meta"] for v in st.session_state.files.values()])
    csv_bytes = meta_df.to_csv(index=False).encode("utf-8")
    st.download_button("📊 Download ALL (CSV metadata)", csv_bytes, file_name="all_metadata.csv")

    # Combined TXT and DOCX
    combined_txt = "".join([f"\n\n===== {fname} =====\n{v['text'] or ''}\n"
                            for fname, v in st.session_state.files.items()])
    st.download_button("🗒️ Download ALL (TXT combined)", combined_txt, file_name="all_cleaned_text.txt")

    if HAS_PYDOCX:
        doc = Document(); doc.add_heading("Cleaned Corpus (All Files)", level=1)
        for fname, v in st.session_state.files.items():
            text, tables_md, figures = v["text"], v["tables"], v["figures"]
            doc.add_heading(fname, level=2)
            if text: doc.add_heading("Text", level=3); doc.add_paragraph(text)
            if tables_md:
                doc.add_heading("Tables (Markdown)", level=3)
                for md in tables_md: doc.add_paragraph(md)
            if figures:
                doc.add_heading("Figures (OCR)", level=3)
                for fg in figures: doc.add_paragraph(f"[page {fg.get('page')}] {fg.get('figure_text','')}")
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        st.download_button("📝 Download ALL (Word/DOCX)", buf, file_name="all_cleaned_corpus.docx")

    # Per-file buttons
    st.markdown("### Per-file downloads")
    for fname, v in st.session_state.files.items():
        col1, col2, col3, col4 = st.columns([2,2,2,2])
        with col1:
            st.download_button(
                label=f"🗒️ TXT — {fname}",
                data=v["text"] or "",
                file_name=f"{os.path.splitext(fname)[0]}_cleaned.txt",
                key=f"txt_{fname}"
            )
        with col2:
            file_jsonl = "\n".join(json.dumps(r, ensure_ascii=False) for r in v["records"])
            st.download_button(
                label=f"📦 JSONL — {fname}",
                data=file_jsonl,
                file_name=f"{os.path.splitext(fname)[0]}_records.jsonl",
                key=f"json_{fname}"
            )
        with col3:
            one_row = pd.DataFrame([v["meta"]]).to_csv(index=False).encode("utf-8")
            st.download_button(
                label=f"📊 CSV — {fname}",
                data=one_row,
                file_name=f"{os.path.splitext(fname)[0]}_meta.csv",
                key=f"csv_{fname}"
            )
        with col4:
            if HAS_PYDOCX:
                d = Document()
                d.add_heading(fname, level=1)
                if v["text"]:
                    d.add_heading("Text", level=2); d.add_paragraph(v["text"])
                if v["tables"]:
                    d.add_heading("Tables (Markdown)", level=2)
                    for md in v["tables"]: d.add_paragraph(md)
                if v["figures"]:
                    d.add_heading("Figures (OCR)", level=2)
                    for fg in v["figures"]: d.add_paragraph(f"[page {fg.get('page')}] {fg.get('figure_text','')}")
                b = io.BytesIO(); d.save(b); b.seek(0)
                st.download_button(
                    label=f"📝 DOCX — {fname}",
                    data=b,
                    file_name=f"{os.path.splitext(fname)[0]}_cleaned.docx",
                    key=f"docx_{fname}"
                )

    with st.expander("Session controls"):
        if st.button("🧹 Clear results (keep selected files)"):
            st.session_state.files = {}
            st.experimental_rerun()

    st.success("Per-file & combined downloads are ready.")
else:
    st.info("Upload files to begin.")
